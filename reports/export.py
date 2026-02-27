"""
Report Export Module

Handles export of reports to various formats:
- PDF (using WeasyPrint or Playwright)
- HTML (static files)
- DOCX (Microsoft Word)
- JSON (structured data)

Each exporter handles conversion, formatting, and file generation.
"""

import json
import os
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

import jinja2


class PDFExporter:
    """
    Export reports to PDF format.
    
    Uses WeasyPrint for PDF generation from HTML/CSS.
    Fallback to Playwright for complex layouts.
    """
    
    def __init__(self, templates_dir: str = "reports/templates"):
        self.templates_dir = Path(templates_dir)
        self.env = jinja2.Environment(
            loader=jinja2.FileSystemLoader(self.templates_dir),
            autoescape=jinja2.select_autoescape(['html', 'xml'])
        )
        
        # Check available PDF engines
        self._check_engines()
    
    def _check_engines(self):
        """Check which PDF generation engines are available."""
        self.weasyprint_available = False
        self.playwright_available = False
        
        try:
            import weasyprint
            self.weasyprint_available = True
        except ImportError:
            pass
        
        try:
            from playwright.sync_api import sync_playwright
            self.playwright_available = True
        except ImportError:
            pass
    
    def export(
        self,
        data: Dict,
        template_name: str,
        output_path: str,
        options: Optional[Dict] = None,
    ) -> str:
        """
        Export report to PDF.
        
        Args:
            data: Report data dictionary
            template_name: Name of Jinja2 template
            output_path: Output PDF file path
            options: Additional export options
            
        Returns:
            Path to generated PDF
        """
        # Render HTML template
        template = self.env.get_template(f"{template_name}.html")
        html_content = template.render(**data)
        
        # Generate PDF
        if self.weasyprint_available:
            return self._export_weasyprint(html_content, output_path, options)
        elif self.playwright_available:
            return self._export_playwright(html_content, output_path, options)
        else:
            raise ImportError(
                "No PDF engine available. Install weasyprint or playwright: "
                "pip install weasyprint playwright"
            )
    
    def _export_weasyprint(
        self,
        html_content: str,
        output_path: str,
        options: Optional[Dict] = None,
    ) -> str:
        """Export using WeasyPrint."""
        from weasyprint import HTML, CSS
        
        # Load CSS if available
        css_content = self._get_default_css()
        css_path = self.templates_dir / "styles.css"
        if css_path.exists():
            css_content = css_path.read_text()
        
        # Generate PDF
        html = HTML(string=html_content)
        css = CSS(string=css_content)
        html.write_pdf(output_path, stylesheets=[css])
        
        return output_path
    
    def _export_playwright(
        self,
        html_content: str,
        output_path: str,
        options: Optional[Dict] = None,
    ) -> str:
        """Export using Playwright (fallback)."""
        from playwright.sync_api import sync_playwright
        
        # Create temp HTML file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False) as f:
            f.write(html_content)
            temp_html = f.name
        
        try:
            with sync_playwright() as p:
                browser = p.chromium.launch()
                page = browser.new_page()
                page.goto(f"file://{temp_html}")
                page.pdf(
                    path=output_path,
                    format="A4",
                    margin={"top": "2cm", "right": "2cm", "bottom": "2cm", "left": "2cm"},
                    print_background=True,
                )
                browser.close()
        finally:
            os.unlink(temp_html)
        
        return output_path
    
    def _get_default_css(self) -> str:
        """Get default CSS styles for PDF."""
        return """
        @page {
            size: A4;
            margin: 2cm;
            @bottom-center {
                content: "Page " counter(page) " of " counter(pages);
                font-size: 9pt;
                color: #666;
            }
        }
        
        body {
            font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
            font-size: 11pt;
            line-height: 1.6;
            color: #333;
        }
        
        h1 {
            font-size: 24pt;
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
            page-break-after: avoid;
        }
        
        h2 {
            font-size: 18pt;
            color: #34495e;
            margin-top: 30px;
            page-break-after: avoid;
        }
        
        h3 {
            font-size: 14pt;
            color: #555;
            page-break-after: avoid;
        }
        
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }
        
        th, td {
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }
        
        th {
            background-color: #f8f9fa;
            font-weight: bold;
        }
        
        .severity-critical {
            background-color: #dc3545;
            color: white;
            padding: 3px 8px;
            border-radius: 3px;
            font-weight: bold;
        }
        
        .severity-high {
            background-color: #fd7e14;
            color: white;
            padding: 3px 8px;
            border-radius: 3px;
            font-weight: bold;
        }
        
        .severity-medium {
            background-color: #ffc107;
            color: #333;
            padding: 3px 8px;
            border-radius: 3px;
            font-weight: bold;
        }
        
        .severity-low {
            background-color: #28a745;
            color: white;
            padding: 3px 8px;
            border-radius: 3px;
        }
        
        .severity-info {
            background-color: #17a2b8;
            color: white;
            padding: 3px 8px;
            border-radius: 3px;
        }
        
        .finding {
            page-break-inside: avoid;
            margin: 20px 0;
            padding: 15px;
            border: 1px solid #ddd;
            border-radius: 5px;
        }
        
        .finding-header {
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 10px;
        }
        
        .finding-title {
            font-size: 14pt;
            font-weight: bold;
            margin: 0;
        }
        
        .evidence-box {
            background-color: #f8f9fa;
            border-left: 4px solid #3498db;
            padding: 10px 15px;
            margin: 10px 0;
            font-family: 'Courier New', monospace;
            font-size: 10pt;
        }
        
        .cover-page {
            text-align: center;
            padding-top: 150px;
        }
        
        .cover-title {
            font-size: 32pt;
            color: #2c3e50;
            margin-bottom: 30px;
        }
        
        .cover-subtitle {
            font-size: 18pt;
            color: #666;
            margin-bottom: 50px;
        }
        
        .cover-meta {
            font-size: 12pt;
            color: #888;
            line-height: 2;
        }
        
        .toc {
            page-break-after: always;
        }
        
        .toc h2 {
            margin-bottom: 20px;
        }
        
        .toc ul {
            list-style: none;
            padding: 0;
        }
        
        .toc li {
            margin: 10px 0;
            display: flex;
            justify-content: space-between;
        }
        
        .risk-meter {
            display: flex;
            justify-content: space-between;
            margin: 20px 0;
            padding: 20px;
            background-color: #f8f9fa;
            border-radius: 5px;
        }
        
        .risk-item {
            text-align: center;
        }
        
        .risk-number {
            font-size: 36pt;
            font-weight: bold;
        }
        
        .risk-label {
            font-size: 10pt;
            color: #666;
            text-transform: uppercase;
        }
        """


class HTMLExporter:
    """
    Export reports to HTML format.
    
    Creates interactive HTML reports with:
    - Search/filter functionality
    - Collapsible sections
    - Responsive design
    - Evidence viewer
    """
    
    def __init__(self, templates_dir: str = "reports/templates"):
        self.templates_dir = Path(templates_dir)
        self.env = jinja2.Environment(
            loader=jinja2.FileSystemLoader(self.templates_dir),
            autoescape=jinja2.select_autoescape(['html', 'xml'])
        )
    
    def export(
        self,
        data: Dict,
        template_name: str,
        output_path: str,
        options: Optional[Dict] = None,
    ) -> str:
        """
        Export report to HTML.
        
        Args:
            data: Report data dictionary
            template_name: Name of Jinja2 template
            output_path: Output HTML file path
            options: Additional export options
            
        Returns:
            Path to generated HTML
        """
        template = self.env.get_template(f"{template_name}.html")
        
        # Add interactive features for HTML
        data['interactive'] = True
        data['export_format'] = 'html'
        
        html_content = template.render(**data)
        
        # Add interactive JavaScript
        if options and options.get('interactive', True):
            html_content = self._add_interactive_features(html_content)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return output_path
    
    def _add_interactive_features(self, html_content: str) -> str:
        """Add JavaScript for interactivity."""
        js_code = """
        <script>
        // Search functionality
        function searchFindings() {
            const query = document.getElementById('searchBox').value.toLowerCase();
            const findings = document.querySelectorAll('.finding');
            findings.forEach(f => {
                const text = f.textContent.toLowerCase();
                f.style.display = text.includes(query) ? 'block' : 'none';
            });
        }
        
        // Filter by severity
        function filterSeverity(severity) {
            const findings = document.querySelectorAll('.finding');
            findings.forEach(f => {
                if (severity === 'all' || f.classList.contains('severity-' + severity)) {
                    f.style.display = 'block';
                } else {
                    f.style.display = 'none';
                }
            });
        }
        
        // Toggle sections
        function toggleSection(id) {
            const el = document.getElementById(id);
            el.style.display = el.style.display === 'none' ? 'block' : 'none';
        }
        
        // Copy to clipboard
        function copyToClipboard(text) {
            navigator.clipboard.writeText(text).then(() => {
                alert('Copied to clipboard!');
            });
        }
        </script>
        """
        
        # Insert before closing body tag
        if '</body>' in html_content:
            html_content = html_content.replace('</body>', js_code + '</body>')
        else:
            html_content += js_code
        
        return html_content


class DOCXExporter:
    """
    Export reports to Microsoft Word format.
    
    Uses python-docx for DOCX generation.
    """
    
    def __init__(self):
        self.available = self._check_availability()
    
    def _check_availability(self) -> bool:
        """Check if python-docx is available."""
        try:
            import docx
            return True
        except ImportError:
            return False
    
    def export(
        self,
        data: Dict,
        output_path: str,
        options: Optional[Dict] = None,
    ) -> str:
        """
        Export report to DOCX.
        
        Args:
            data: Report data dictionary
            output_path: Output DOCX file path
            options: Additional export options
            
        Returns:
            Path to generated DOCX
        """
        if not self.available:
            raise ImportError(
                "python-docx required. Install with: pip install python-docx"
            )
        
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title page
        title = doc.add_heading(data.get('title', 'Penetration Test Report'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if data.get('company_name'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(data['company_name'])
            run.font.size = Pt(16)
            run.font.bold = True
        
        if data.get('report_date'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"Report Date: {data['report_date']}")
        
        doc.add_page_break()
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(data.get('executive_summary', 'No summary available.'))
        
        # Risk Summary
        doc.add_heading('Risk Summary', level=1)
        risk_data = data.get('risk_summary', {})
        
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        severity_rows = [
            ('Critical', risk_data.get('critical', 0)),
            ('High', risk_data.get('high', 0)),
            ('Medium', risk_data.get('medium', 0)),
            ('Low', risk_data.get('low', 0)),
            ('Info', risk_data.get('info', 0)),
        ]
        
        for i, (severity, count) in enumerate(severity_rows):
            table.rows[i].cells[0].text = severity
            table.rows[i].cells[1].text = str(count)
        
        # Findings
        doc.add_heading('Detailed Findings', level=1)
        
        for finding in data.get('findings', []):
            doc.add_heading(finding.get('title', 'Untitled'), level=2)
            
            p = doc.add_paragraph()
            p.add_run(f"Severity: ").bold = True
            p.add_run(finding.get('severity', 'Unknown'))
            
            doc.add_heading('Description', level=3)
            doc.add_paragraph(finding.get('description', ''))
            
            if finding.get('proof_of_concept'):
                doc.add_heading('Proof of Concept', level=3)
                doc.add_paragraph(finding['proof_of_concept'])
            
            if finding.get('remediation'):
                doc.add_heading('Remediation', level=3)
                doc.add_paragraph(finding['remediation'])
            
            doc.add_paragraph()  # Spacing
        
        # Save document
        doc.save(output_path)
        
        return output_path


class JSONExporter:
    """
    Export reports to JSON format.
    
    Structured data export for integration with other tools.
    """
    
    def export(
        self,
        data: Dict,
        output_path: str,
        options: Optional[Dict] = None,
    ) -> str:
        """
        Export report to JSON.
        
        Args:
            data: Report data dictionary
            output_path: Output JSON file path
            options: Additional export options (indent, etc.)
            
        Returns:
            Path to generated JSON
        """
        indent = options.get('indent', 2) if options else 2
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=indent, default=str, ensure_ascii=False)
        
        return output_path
    
    def export_to_string(self, data: Dict, indent: int = 2) -> str:
        """Export report to JSON string."""
        return json.dumps(data, indent=indent, default=str, ensure_ascii=False)
